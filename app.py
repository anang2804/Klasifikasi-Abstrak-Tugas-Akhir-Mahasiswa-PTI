"""
Flask application untuk klasifikasi abstrak tugas akhir
"""
import os
import json
from flask import Flask, render_template, request, jsonify, flash, redirect, url_for
from werkzeug.utils import secure_filename
import pandas as pd

from config import Config
from models import db, Abstract, ModelMetrics, ClassificationHistory
from scraper import scrape_and_save
from classifier import KNNClassifier


app = Flask(__name__)
app.config.from_object(Config)

# Initialize database
db.init_app(app)

# Global classifier instance
classifier = None


def init_classifier():
    """Initialize atau load classifier"""
    global classifier
    
    if classifier is None:
        classifier = KNNClassifier(k=app.config['KNN_K_VALUE'])
        
        # Coba load model yang sudah ada
        if os.path.exists('models/knn_classifier.joblib'):
            try:
                classifier.load('models')
                print(f"✅ Model loaded successfully! is_trained={classifier.is_trained}")
            except Exception as e:
                print(f"❌ Error loading model: {e}")
                classifier = None
        else:
            print("⚠️ Model file not found: models/knn_classifier.joblib")
            print("   Please train the model first!")
    else:
        print(f"ℹ️ Classifier already initialized. is_trained={classifier.is_trained}")


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def extract_abstract_section(text):
    """Extract hanya bagian ABSTRAK dari teks lengkap"""
    import re
    
    # Pattern untuk mencari bagian abstrak (case-insensitive)
    patterns = [
        # Pattern 1: ABSTRAK ... (KATA KUNCI|BAB|ABSTRACT|PENDAHULUAN|DAFTAR)
        r'(?:ABSTRAK|Abstrak)\s*\n(.*?)(?:\n\s*(?:KATA KUNCI|BAB\s+[IVX]|ABSTRACT|PENDAHULUAN|DAFTAR|CHAPTER))',
        # Pattern 2: ABSTRACT ... (KEYWORDS|CHAPTER|BAB|ABSTRAK)
        r'(?:ABSTRACT|Abstract)\s*\n(.*?)(?:\n\s*(?:KEYWORDS|CHAPTER|BAB\s+[IVX]|ABSTRAK|PENDAHULUAN))',
        # Pattern 3: Lebih fleksibel - ambil 1000 karakter pertama setelah kata ABSTRAK
        r'(?:ABSTRAK|Abstrak|ABSTRACT|Abstract)\s*[:\n](.*)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            abstract = match.group(1).strip()
            # Batasi maksimal 3000 karakter (abstrak biasanya 150-500 kata)
            if len(abstract) > 3000:
                abstract = abstract[:3000]
            # Minimal 100 karakter untuk dianggap valid
            if len(abstract) >= 100:
                return abstract
    
    # Fallback: jika tidak ketemu pattern, ambil 1500 karakter pertama
    # (asumsi abstrak ada di awal dokumen)
    return text[:1500].strip()


def extract_text_from_file(filepath):
    """Extract text from uploaded file - FOKUS KE BAGIAN ABSTRAK"""
    ext = filepath.rsplit('.', 1)[1].lower()
    
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                full_text = f.read()
                # Coba ekstrak bagian abstrak
                return extract_abstract_section(full_text)
        
        elif ext == 'pdf':
            # Extract dari PDF menggunakan PyPDF2
            import PyPDF2
            text = ""
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                # Abstrak biasanya di 3-10 halaman pertama
                # (setelah cover, pengesahan, kata pengantar)
                max_pages = min(len(pdf_reader.pages), 15)
                for i in range(max_pages):
                    page_text = pdf_reader.pages[i].extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                    # Cek apakah sudah ketemu bagian abstrak
                    if 'ABSTRAK' in text.upper() or 'ABSTRACT' in text.upper():
                        # Cukup baca sampai 5 halaman setelah ketemu abstrak
                        if i >= 3:
                            break
            
            # Ekstrak hanya bagian abstrak
            return extract_abstract_section(text)
        
        elif ext == 'docx':
            # Extract dari DOCX menggunakan python-docx
            from docx import Document
            doc = Document(filepath)
            text = ""
            found_abstract = False
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                
                # Cek apakah ini heading ABSTRAK/ABSTRACT
                if para_text.upper() in ['ABSTRAK', 'ABSTRACT']:
                    found_abstract = True
                    continue
                
                # Jika sudah ketemu abstrak, ambil teksnya
                if found_abstract:
                    text += para_text + "\n"
                    
                    # Stop jika ketemu heading lain (BAB, KATA KUNCI, dll)
                    if any(keyword in para_text.upper() for keyword in 
                           ['KATA KUNCI', 'BAB I', 'BAB 1', 'PENDAHULUAN', 'KEYWORDS', 'CHAPTER']):
                        break
                    
                    # Atau jika sudah cukup panjang (3000 karakter)
                    if len(text) > 3000:
                        break
            
            # Jika tidak ketemu dengan cara di atas, coba pattern matching
            if not text or len(text) < 100:
                full_text = "\n".join([p.text for p in doc.paragraphs])
                return extract_abstract_section(full_text)
            
            return text.strip()
        
    except Exception as e:
        print(f"Error extracting text from {filepath}: {str(e)}")
        return ""
    
    return ""


@app.route('/')
def index():
    """Halaman utama - menampilkan daftar abstrak"""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Query abstracts dengan pagination
    abstracts_query = Abstract.query.order_by(Abstract.year.desc(), Abstract.id.desc())
    pagination = abstracts_query.paginate(page=page, per_page=per_page, error_out=False)
    abstracts = pagination.items
    
    # Statistik
    total_abstracts = Abstract.query.count()
    # Hitung dari label manual (bukan predicted_label)
    total_rpl = Abstract.query.filter_by(label='RPL').count()
    total_tkj = Abstract.query.filter_by(label='TKJ').count()
    
    # Total labeled: hitung yang punya label manual
    total_labeled = Abstract.query.filter(Abstract.label.isnot(None)).count()
    
    stats = {
        'total': total_abstracts,
        'rpl': total_rpl,
        'tkj': total_tkj,
        'labeled': total_labeled,
        'unlabeled': total_abstracts - total_labeled
    }
    
    return render_template('index.html', 
                         abstracts=abstracts, 
                         pagination=pagination,
                         stats=stats)


@app.route('/data-test')
def data_test():
    """Halaman untuk menampilkan data uji (hasil klasifikasi)"""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Query classification history dengan pagination
    history_query = ClassificationHistory.query.order_by(ClassificationHistory.classified_at.desc())
    pagination = history_query.paginate(page=page, per_page=per_page, error_out=False)
    test_data = pagination.items
    
    # Statistik
    total_test = ClassificationHistory.query.count()
    total_rpl = ClassificationHistory.query.filter_by(predicted_label='RPL').count()
    total_tkj = ClassificationHistory.query.filter_by(predicted_label='TKJ').count()
    total_manual = ClassificationHistory.query.filter_by(source='manual').count()
    total_upload = ClassificationHistory.query.filter_by(source='upload').count()
    
    stats = {
        'total': total_test,
        'rpl': total_rpl,
        'tkj': total_tkj,
        'manual': total_manual,
        'upload': total_upload
    }
    
    return render_template('test_data.html',
                         test_data=test_data,
                         pagination=pagination,
                         stats=stats)


@app.route('/data-test/delete/<int:id>', methods=['POST'])
def delete_test_data(id):
    """Hapus data uji berdasarkan ID"""
    try:
        test_data = ClassificationHistory.query.get_or_404(id)
        db.session.delete(test_data)
        db.session.commit()
        flash('Data uji berhasil dihapus!', 'success')
    except Exception as e:
        flash(f'Error menghapus data: {str(e)}', 'error')
    
    return redirect(url_for('data_test'))


@app.route('/data-test/clear-all', methods=['POST'])
def clear_all_test_data():
    """Hapus semua data uji"""
    try:
        count = ClassificationHistory.query.count()
        ClassificationHistory.query.delete()
        db.session.commit()
        flash(f'✅ Berhasil menghapus {count} data uji!', 'success')
    except Exception as e:
        flash(f'Error menghapus data: {str(e)}', 'error')
    
    return redirect(url_for('data_test'))


@app.route('/scrape', methods=['GET', 'POST'])
def scrape():
    """Halaman untuk scraping data"""
    if request.method == 'POST':
        start_year = request.form.get('start_year', app.config['START_YEAR'], type=int)
        end_year = request.form.get('end_year', app.config['END_YEAR'], type=int)
        
        try:
            # Scraping dengan auto-label otomatis menggunakan keyword scoring
            result = scrape_and_save(
                app.config['BASE_URL'],
                start_year,
                end_year,
                auto_label=True  # Otomatis label dengan keyword scoring
            )
            
            flash(result['message'], 'success')
            
            # Redirect ke halaman index untuk melihat hasil
            if result.get('total_saved', 0) > 0:
                if result.get('auto_labeled', 0) > 0:
                    flash(f'✅ Data otomatis di-label: RPL={result.get("rpl_count", 0)}, TKJ={result.get("tkj_count", 0)}. Anda bisa koreksi label di halaman Label Data jika diperlukan.', 'info')
                return redirect(url_for('index'))
            else:
                return redirect(url_for('index'))
            
        except Exception as e:
            flash(f'Error during scraping: {str(e)}', 'error')
    
    return render_template('scrape.html')


@app.route('/label')
def label_data():
    """Halaman untuk menampilkan data latih (training data yang sudah dilabel)"""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Query data yang sudah dilabel (data latih)
    training_query = Abstract.query.filter(Abstract.label.isnot(None)).order_by(Abstract.year.desc(), Abstract.id.desc())
    pagination = training_query.paginate(page=page, per_page=per_page, error_out=False)
    training_data = pagination.items
    
    # Statistik
    total_training = Abstract.query.filter(Abstract.label.isnot(None)).count()
    total_rpl = Abstract.query.filter_by(label='RPL').count()
    total_tkj = Abstract.query.filter_by(label='TKJ').count()
    
    stats = {
        'total': total_training,
        'rpl': total_rpl,
        'tkj': total_tkj
    }
    
    return render_template('label_data.html', 
                         training_data=training_data,
                         pagination=pagination,
                         stats=stats)


@app.route('/api/label/<int:abstract_id>', methods=['POST'])
def api_label_abstract(abstract_id):
    """API untuk memberikan label pada abstract"""
    data = request.get_json()
    label = data.get('label')
    
    if label not in ['RPL', 'TKJ']:
        return jsonify({'error': 'Invalid label'}), 400
    
    abstract = Abstract.query.get_or_404(abstract_id)
    abstract.label = label
    abstract.is_training_data = True
    
    try:
        db.session.commit()
        return jsonify({'success': True, 'message': 'Label saved'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auto-label-unlabeled', methods=['POST'])
def api_auto_label_unlabeled():
    """API untuk auto-label semua data yang belum berlabel"""
    global classifier
    
    if classifier is None or not classifier.is_trained:
        return jsonify({'error': 'Model belum di-train'}), 400
    
    try:
        # Ambil semua data yang benar-benar belum berlabel (label IS NULL AND predicted_label IS NULL)
        unlabeled = Abstract.query.filter(
            Abstract.label.is_(None),
            Abstract.predicted_label.is_(None)
        ).all()
        
        if not unlabeled:
            return jsonify({
                'success': True,
                'labeled': 0,
                'message': 'Tidak ada data yang perlu dilabel'
            })
        
        print(f"\n🤖 Auto-labeling {len(unlabeled)} unlabeled data...")
        
        # Klasifikasi batch
        texts = [abstract.abstract_text for abstract in unlabeled]
        predictions = classifier.predict(texts)
        probabilities = classifier.predict_proba(texts)
        
        # Update database - set predicted_label
        for i, abstract in enumerate(unlabeled):
            abstract.predicted_label = predictions[i]
            abstract.confidence = float(probabilities[i].max())
        
        db.session.commit()
        
        # Hitung distribusi label
        rpl_count = sum(1 for p in predictions if p == 'RPL')
        tkj_count = sum(1 for p in predictions if p == 'TKJ')
        
        print(f"✓ Auto-labeling complete! RPL: {rpl_count}, TKJ: {tkj_count}")
        
        return jsonify({
            'success': True,
            'labeled': len(unlabeled),
            'rpl_count': rpl_count,
            'tkj_count': tkj_count,
            'message': f'Successfully labeled {len(unlabeled)} abstracts'
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"❌ Error during auto-labeling: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/train', methods=['GET', 'POST'])
def train_model():
    """Halaman dan API untuk training model"""
    global classifier
    
    if request.method == 'POST':
        try:
            # ✅ HANYA AMBIL DATA LABEL MANUAL (Best Practice)
            training_data = Abstract.query.filter(
                Abstract.label.isnot(None)
            ).all()
            
            if len(training_data) < 10:
                flash('Minimal 10 data training dengan label manual diperlukan!', 'error')
                return redirect(url_for('train_model'))
            
            # Ekstrak texts dan labels (hanya manual labels)
            texts = [abstract.abstract_text for abstract in training_data]
            labels = [abstract.label for abstract in training_data]
            
            print(f"\n{'='*60}")
            print(f"🎯 TRAINING WITH MANUAL LABELS ONLY")
            print(f"{'='*60}")
            print(f"Total training data: {len(training_data)}")
            
            # Get K value dari form
            k_value = request.form.get('k_value', 5, type=int)
            
            # Initialize classifier
            classifier = KNNClassifier(k=k_value)
            
            # Prepare data with STRATIFIED split (sudah ada di classifier.py)
            data = classifier.prepare_data(
                texts, labels,
                test_size=app.config['TEST_SIZE'],
                random_state=app.config['RANDOM_STATE']
            )
            
            # Train
            classifier.train(data['X_train'], data['y_train'])
            
            # Evaluate
            evaluation = classifier.evaluate(data['X_test'], data['y_test'])
            
            # Save model
            classifier.save('models')
            
            # Save metrics to database
            metrics = ModelMetrics(
                k_value=k_value,
                accuracy=evaluation['accuracy'],
                precision_rpl=evaluation['precision'].get('RPL', 0),
                precision_tkj=evaluation['precision'].get('TKJ', 0),
                recall_rpl=evaluation['recall'].get('RPL', 0),
                recall_tkj=evaluation['recall'].get('TKJ', 0),
                f1_rpl=evaluation['f1_score'].get('RPL', 0),
                f1_tkj=evaluation['f1_score'].get('TKJ', 0),
                training_samples=len(data['y_train']),
                test_samples=len(data['y_test'])
            )
            db.session.add(metrics)
            db.session.commit()
            
            flash(f'✅ Model trained successfully! Accuracy: {evaluation["accuracy"]:.2%}', 'success')
            
            return redirect(url_for('evaluation'))
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'Error during training: {str(e)}', 'error')
    
    # GET request - hitung statistik data
    # ✅ Data Latih (dari scraping + auto-label)
    training_count = Abstract.query.filter(Abstract.label.isnot(None)).count()
    
    # ✅ Data Uji (dari klasifikasi manual/upload)
    test_count = ClassificationHistory.query.count()
    
    # Total semua data
    total_data = Abstract.query.count()
    
    return render_template('train.html', 
                         training_count=training_count,
                         test_count=test_count,
                         total_data=total_data)


@app.route('/classify', methods=['GET', 'POST'])
def classify():
    """Halaman untuk klasifikasi abstrak baru (Data Uji)"""
    global classifier
    
    if classifier is None or not classifier.is_trained:
        flash('Model belum di-train! Silakan train model terlebih dahulu.', 'warning')
        return redirect(url_for('train_model'))
    
    result = None
    
    if request.method == 'POST':
        text = request.form.get('abstract_text', '')
        
        if text:
            try:
                predicted_label, confidence = classifier.predict_single(text)
                
                # Simpan ke ClassificationHistory sebagai Data Uji
                history = ClassificationHistory(
                    abstract_text=text,
                    predicted_label=predicted_label,
                    confidence=confidence,
                    source='manual'
                )
                db.session.add(history)
                db.session.commit()
                
                # Get important words for highlighting (increased to 20 for better coverage)
                important_words = classifier.get_important_words(text, top_n=20)
                
                # Collect all keyword variations for JavaScript highlighting
                all_variations = []
                for word_info in important_words['original_words']:
                    all_variations.extend(word_info['variations'])
                
                result = {
                    'text': text,
                    'label': predicted_label,
                    'confidence': confidence,
                    'important_words': important_words['words'],
                    'original_words': important_words['original_words'],
                    'preprocessed': important_words['preprocessed_text'],
                    'keyword_variations': all_variations  # For JavaScript highlighting
                }
            except Exception as e:
                flash(f'Error during classification: {str(e)}', 'error')
    
    return render_template('classify.html', result=result)


@app.route('/upload', methods=['POST'])
def upload_file():
    """Upload dan klasifikasi file"""
    global classifier
    
    if classifier is None or not classifier.is_trained:
        return jsonify({'error': 'Model belum di-train'}), 400
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(filepath)
            
            # Extract text
            text = extract_text_from_file(filepath)
            
            if not text:
                return jsonify({'error': 'Could not extract text from file'}), 400
            
            # Classify
            predicted_label, confidence = classifier.predict_single(text)
            
            # Simpan ke ClassificationHistory sebagai Data Uji
            history = ClassificationHistory(
                abstract_text=text,
                predicted_label=predicted_label,
                confidence=confidence,
                source='upload'
            )
            db.session.add(history)
            db.session.commit()
            
            # Clean up
            os.remove(filepath)
            
            return jsonify({
                'success': True,
                'label': predicted_label,
                'confidence': float(confidence)
            })
            
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400


@app.route('/classify-all', methods=['POST'])
def classify_all():
    """Klasifikasi semua abstrak yang belum diklasifikasi"""
    global classifier
    
    if classifier is None or not classifier.is_trained:
        return jsonify({'error': 'Model belum di-train'}), 400
    
    try:
        # Ambil abstrak yang belum diklasifikasi
        abstracts = Abstract.query.filter(Abstract.predicted_label.is_(None)).all()
        
        if not abstracts:
            return jsonify({'message': 'Tidak ada abstrak yang perlu diklasifikasi'})
        
        # Klasifikasi batch
        texts = [abstract.abstract_text for abstract in abstracts]
        predictions = classifier.predict(texts)
        probabilities = classifier.predict_proba(texts)
        
        # Update database
        for i, abstract in enumerate(abstracts):
            abstract.predicted_label = predictions[i]
            abstract.confidence = float(probabilities[i].max())
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'classified': len(abstracts),
            'message': f'Successfully classified {len(abstracts)} abstracts'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/evaluation')
def evaluation():
    """Halaman evaluasi model"""
    # Ambil metrics terbaru
    latest_metrics = ModelMetrics.query.order_by(ModelMetrics.trained_at.desc()).first()
    
    # Ambil semua metrics untuk grafik perbandingan
    all_metrics = ModelMetrics.query.order_by(ModelMetrics.trained_at.desc()).limit(10).all()
    
    return render_template('evaluation.html', 
                         latest=latest_metrics,
                         history=all_metrics)


@app.route('/api/stats')
def api_stats():
    """API untuk statistik"""
    total_abstracts = Abstract.query.count()
    total_rpl = Abstract.query.filter_by(predicted_label='RPL').count()
    total_tkj = Abstract.query.filter_by(predicted_label='TKJ').count()
    
    # Distribusi per tahun
    year_distribution = db.session.query(
        Abstract.year,
        Abstract.predicted_label,
        db.func.count(Abstract.id)
    ).filter(Abstract.predicted_label.isnot(None))\
     .group_by(Abstract.year, Abstract.predicted_label)\
     .all()
    
    return jsonify({
        'total': total_abstracts,
        'rpl': total_rpl,
        'tkj': total_tkj,
        'year_distribution': [
            {'year': year, 'label': label, 'count': count}
            for year, label, count in year_distribution
        ]
    })


@app.route('/abstract/<int:abstract_id>')
def view_abstract(abstract_id):
    """View detail abstrak"""
    abstract = Abstract.query.get_or_404(abstract_id)
    return render_template('abstract_detail.html', abstract=abstract)


@app.before_request
def before_first_request():
    """Initialize database dan classifier sebelum request pertama"""
    db.create_all()
    init_classifier()


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        init_classifier()
    
    app.run(debug=True, host='0.0.0.0', port=5000)
